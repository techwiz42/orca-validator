"""Render Markdown to a .docx (python-docx). Lightweight — headings, bullets, **bold** runs;
everything else as paragraphs. Good enough for a redlined-contract export the user can edit."""
import io
import re

from docx import Document


def markdown_to_docx(md: str, title: str = "Revised Document") -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            doc.add_heading(heading.group(2), level=min(len(heading.group(1)), 4))
            continue
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^[-*]\s+", "", line))
            continue
        _add_runs(doc.add_paragraph(), line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs(paragraph, text: str) -> None:
    # segments inside **...** become bold runs
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True
