"""Document → text, dispatched by format.

Ingests most common document formats (the service is *tuned* for contracts via the downstream
machine, but not limited to one input format):
  - PDF / images / EPUB / XPS / …  → PyMuPDF, with a tesseract OCR fallback for image-only pages
  - .docx                          → python-docx
  - plain text (.txt/.md/.csv/…)   → read directly

Returns an OCRResult with per-source provenance. On undecodable input it sets `error` (the caller
fails the run loudly rather than passing an empty field set downstream).
"""
import io
import logging
import os

from backend.ocr.result import OCRResult

logger = logging.getLogger(__name__)

_MIN_EMBEDDED_CHARS = 20  # below this, a PDF page is treated as image-only and sent to OCR
_FITZ_EXTS = {"pdf", "png", "jpg", "jpeg", "tif", "tiff", "bmp", "gif", "webp",
              "epub", "xps", "fb2", "cbz", "svg"}
_TEXT_EXTS = {"txt", "md", "markdown", "text", "csv", "log", "rtf"}

# Abuse / decompression-bomb guards.
MAX_PAGES = 100               # reject documents with more than this many pages
MAX_RENDER_PIXELS = 25_000_000  # cap OCR-render size so a giant page can't OOM the parser


def read_document(file_path: str, dpi: int = 200) -> OCRResult:
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    try:
        if ext == "docx":
            return _read_docx(file_path)
        if ext in _TEXT_EXTS:
            return _read_text(file_path)
        if ext in _FITZ_EXTS or ext == "":
            return _read_fitz(file_path, dpi)
        return OCRResult(error=f"unsupported document type: .{ext}")
    except Exception as e:  # noqa: BLE001 — fail loudly with the reason, never a fake empty result
        return OCRResult(error=f"could not read document (.{ext or 'unknown'}): {e}")


def _read_text(file_path: str) -> OCRResult:
    with open(file_path, encoding="utf-8", errors="replace") as f:
        raw = f.read().strip()
    if not raw:
        return OCRResult(error="empty text document")
    return OCRResult(raw_text=raw, pages=[{"page": 0, "method": "text", "chars": len(raw)}])


def _read_docx(file_path: str) -> OCRResult:
    try:
        import docx  # python-docx
    except ImportError as e:
        return OCRResult(error=f"python-docx not available: {e}")
    d = docx.Document(file_path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    raw = "\n".join(parts).strip()
    if not raw:
        return OCRResult(error="no text extracted from .docx")
    return OCRResult(raw_text=raw, pages=[{"page": 0, "method": "docx", "chars": len(raw)}])


def _read_fitz(file_path: str, dpi: int) -> OCRResult:
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        return OCRResult(error=f"PyMuPDF (fitz) not available: {e}")

    doc = fitz.open(file_path)
    if doc.page_count > MAX_PAGES:
        doc.close()
        return OCRResult(error=f"document exceeds the {MAX_PAGES}-page limit ({doc.page_count} pages)")
    pages: list[dict] = []
    texts: list[str] = []
    try:
        for i, page in enumerate(doc):
            text = page.get_text() or ""
            method = "embedded"
            if len(text.strip()) < _MIN_EMBEDDED_CHARS:
                ocr_text = _ocr_page_image(page, dpi)
                if ocr_text is not None:
                    text, method = ocr_text, "tesseract"
            pages.append({"page": i, "method": method, "chars": len(text)})
            texts.append(text)
    finally:
        doc.close()

    raw = "\n\n".join(texts).strip()
    if not raw:
        return OCRResult(pages=pages, error="no text could be extracted (embedded or OCR)")
    return OCRResult(raw_text=raw, pages=pages)


def _ocr_page_image(page, dpi: int) -> str | None:
    """Render a page to PNG and OCR it with tesseract. None if tesseract is unavailable.

    Caps the render resolution (image-bomb guard) so an enormous page can't allocate a
    multi-gigabyte pixmap and OOM the parser.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return None
    # Decompression-bomb guard for PIL, and a fitz-render-size cap computed BEFORE allocating.
    Image.MAX_IMAGE_PIXELS = MAX_RENDER_PIXELS
    rect = page.rect
    est_px = (rect.width * dpi / 72.0) * (rect.height * dpi / 72.0)
    if est_px > MAX_RENDER_PIXELS and est_px > 0:
        dpi = max(36, int(dpi * (MAX_RENDER_PIXELS / est_px) ** 0.5))
    try:
        pix = page.get_pixmap(dpi=dpi)
        return pytesseract.image_to_string(Image.open(io.BytesIO(pix.tobytes("png"))))
    except Exception as e:  # noqa: BLE001 — includes PIL DecompressionBombError
        logger.warning("tesseract OCR failed/blocked for a page: %s", e)
        return None
